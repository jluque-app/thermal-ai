# ThermalAI_report.py
from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================
# BRANDING / ASSETS
# ============================
ASSETS_DIR = Path("sources")
LOGO_PATH = ASSETS_DIR / "logo3.jpg"   # required branding
EER_PATH = ASSETS_DIR / "EER.png"      # optional illustrative graphic

# Legacy-ish colors
COLOR_MAROON = "7C3A32"
COLOR_ROW_A = "D9D9D9"
COLOR_ROW_B = "EFEFEF"

# Page geometry (A4; python-docx uses inches)
MARGIN_TOP = Inches(0.10)
MARGIN_BOTTOM = Inches(0.20)
MARGIN_LEFT = Inches(0.55)
MARGIN_RIGHT = Inches(0.55)

# Typography
FONT_BODY = "Calibri"
FONT_BRAND = "Cambria"
BRAND_SPACED = "T H E R M A L A I"
SUPPORT_EMAIL = "info@allretech.org"  # edit if needed

# EPC-adjacent but legally safe disclaimer
DISCLAIMER_FOOTER = (
    "DISCLAIMER: This report is an engineering screening based on thermal imaging and AI-driven image analysis. "
    "It does not constitute an official Energy Performance Certificate (EPC) nor a regulated assessment. "
    "All outputs are indicative estimates and may be subject to significant uncertainty due to measurement conditions, "
    "assumptions, and model limitations. No responsibility is accepted for decisions made based on this report."
)


# ============================
# BASIC HELPERS
# ============================
def _safe(v, default="—"):
    return default if v is None else v


def _num(v: Any, decimals: int = 2) -> Optional[float]:
    try:
        if v is None:
            return None
        return round(float(v), decimals)
    except Exception:
        return None


def _fmt(v: Any, unit: str = "", decimals: int = 2) -> str:
    x = _num(v, decimals)
    if x is None:
        return "—"
    if decimals == 0:
        s = f"{int(round(x))}"
    else:
        s = f"{x:.{decimals}f}"
    return f"{s}{unit}"


def _set_font_run(run, name: str, size_pt: int, bold: bool = False, italic: bool = False, underline: bool = False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def _set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_bold(cell, bold=True) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold


def _add_field(run, field_code: str) -> None:
    """
    Insert a Word field (PAGE, NUMPAGES) into a run using w:fldChar elements.
    """
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def _setup_document(doc: Document) -> None:
    # Apply margins to all existing sections
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10)


# ============================
# HEADER / FOOTER
# ============================
def _set_cover_footer(section) -> None:
    """
    Cover page footer: brand + email only; no page numbering.
    """
    footer = section.footer
    footer.paragraphs[0].text = ""

    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p1.add_run(BRAND_SPACED)
    _set_font_run(r1, FONT_BRAND, 9)

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = p2.add_run(SUPPORT_EMAIL)
    _set_font_run(r2, FONT_BRAND, 9, underline=True)


def _set_main_header_footer(section, disclaimer_text: str) -> None:
    """
    Main report header: logo top-right.
    Main report footer: brand centered, disclaimer, and right-aligned page numbering.
    """
    header = section.header
    header.paragraphs[0].text = ""
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run()
    if LOGO_PATH.exists():
        try:
            hr.add_picture(str(LOGO_PATH), width=Inches(1.10))
        except Exception:
            pass

    footer = section.footer
    footer.paragraphs[0].text = ""

    p_brand = footer.paragraphs[0]
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_brand = p_brand.add_run(BRAND_SPACED)
    _set_font_run(r_brand, FONT_BRAND, 9)

    footer.add_paragraph("")  # spacer

    p_disc = footer.add_paragraph()
    p_disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_disc = p_disc.add_run(disclaimer_text)
    _set_font_run(r_disc, FONT_BODY, 7, italic=True)

    footer.add_paragraph("")  # spacer

    p_page = footer.add_paragraph()
    p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r0 = p_page.add_run("P a g e    ")
    _set_font_run(r0, FONT_BRAND, 9)

    r_page = p_page.add_run()
    _add_field(r_page, "PAGE")
    _set_font_run(r_page, FONT_BRAND, 9)

    r_mid = p_page.add_run(" | ")
    _set_font_run(r_mid, FONT_BRAND, 9)

    r_total = p_page.add_run()
    _add_field(r_total, "NUMPAGES")
    _set_font_run(r_total, FONT_BRAND, 9)


# ============================
# LAYOUT BLOCKS
# ============================
def _add_cover_page(doc: Document, address: str, created_at: str, analysis_id: str) -> None:
    """
    Section 0: Cover
    """
    sec0 = doc.sections[0]
    _set_cover_footer(sec0)

    # Use a 1-col table to control vertical spacing (legacy look)
    t = doc.add_table(rows=7, cols=1)
    t.autofit = True

    # Spacer
    t.cell(0, 0).text = ""

    # Title
    p = t.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("THERMALAI")
    _set_font_run(r, FONT_BODY, 20, bold=True)

    # Subtitle (safe wording)
    p2 = t.cell(2, 0).paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("ENGINEERING SCREENING REPORT\n(AI + THERMAL IMAGING)")
    _set_font_run(r2, FONT_BODY, 12, bold=True)

    # Spacer
    t.cell(3, 0).text = ""

    # Address
    p4 = t.cell(4, 0).paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(address or "—")
    _set_font_run(r4, FONT_BODY, 12)

    # Meta
    p5 = t.cell(5, 0).paragraphs[0]
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run(f"Created: {created_at or '—'}    |    Analysis ID: {analysis_id}")
    _set_font_run(r5, FONT_BODY, 10)

    # Spacer
    t.cell(6, 0).text = ""

    doc.add_page_break()


def _start_main_section(doc: Document) -> None:
    """
    New section starting on a new page so header/footer differs from cover.
    """
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.top_margin = MARGIN_TOP
    sec.bottom_margin = MARGIN_BOTTOM
    sec.left_margin = MARGIN_LEFT
    sec.right_margin = MARGIN_RIGHT
    _set_main_header_footer(sec, DISCLAIMER_FOOTER)


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font_run(r, FONT_BODY, 14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font_run(r, FONT_BODY, 12, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_par(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font_run(r, FONT_BODY, 10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_legacy_style_table(doc: Document, title_left: str, title_right: str, rows: list[tuple[str, str]]) -> None:
    """
    Legacy style: maroon header row + zebra body; left column bold.
    """
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = title_left
    hdr[1].text = title_right

    _set_cell_shading(hdr[0], COLOR_MAROON)
    _set_cell_shading(hdr[1], COLOR_MAROON)
    _set_cell_bold(hdr[0], True)
    _set_cell_bold(hdr[1], True)

    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                _set_font_run(r, FONT_BODY, 12, bold=True)

    # Body rows
    for i, (label, val) in enumerate(rows):
        tr = table.add_row().cells
        tr[0].text = label
        tr[1].text = val

        fill = COLOR_ROW_A if (i % 2 == 0) else COLOR_ROW_B
        _set_cell_shading(tr[0], fill)
        _set_cell_shading(tr[1], fill)

        _set_cell_bold(tr[0], True)

        for c in tr:
            for p in c.paragraphs:
                for r in p.runs:
                    _set_font_run(r, FONT_BODY, 10, bold=r.bold)

    doc.add_paragraph("")  # spacing


def _try_decode_b64(b64_str: Any) -> Optional[bytes]:
    if not isinstance(b64_str, str):
        return None
    s = b64_str.strip()
    if not s:
        return None
    try:
        return base64.b64decode(s)
    except Exception:
        return None


def _extract_images_from_payload(raw: Dict[str, Any]) -> Tuple[Optional[bytes], Optional[bytes]]:
    """
    Returns (rgb_bytes, overlay_bytes) if available.
    Non-breaking: does not require backend changes; tries multiple likely keys.
    """
    artifacts = (raw.get("artifacts") or {})

    overlay = _try_decode_b64(artifacts.get("overlay_image_base64_png"))

    rgb = (
        _try_decode_b64(artifacts.get("rgb_image_base64_png"))
        or _try_decode_b64(artifacts.get("rgb_image_base64_jpg"))
        or _try_decode_b64(artifacts.get("rgb_image_base64"))
        or _try_decode_b64(artifacts.get("rgb_base64_png"))
        or _try_decode_b64(artifacts.get("input_rgb_base64_png"))
    )

    return rgb, overlay


def _add_figure_page_from_bytes(doc: Document, title: str, image_bytes: bytes, width_in: float = 6.6) -> None:
    doc.add_page_break()
    _add_h2(doc, title)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    try:
        r.add_picture(BytesIO(image_bytes), width=Inches(width_in))
    except Exception:
        _add_par(doc, "Image was provided but could not be rendered.")


# ============================
# PUBLIC API: used by docx_endpoint.py
# ============================
def build_docx_bytes(payload: Dict[str, Any]) -> bytes:
    """
    Drop-in entry point used by /v1/report/docx.
    Expects:
      payload = {"report": {...}, "raw": {...}}
    """
    report = payload.get("report") or {}
    raw = payload.get("raw") or {}

    meta = report.get("meta", {}) or {}
    headline = report.get("headline", {}) or {}
    assumptions = report.get("assumptions", []) or []

    analysis_id = str(meta.get("analysis_id") or raw.get("analysis_id") or "analysis")

    inputs = (raw.get("inputs") or {})
    results = (raw.get("results") or {})
    comps = (results.get("components") or {})
    totals = (results.get("totals") or {})

    # Cover metadata
    city = str(_safe(meta.get("city"), "—"))
    created_at = str(_safe(meta.get("created_at"), "—"))

    # Try to find an address-like value without breaking schema
    address = meta.get("address") or inputs.get("address") or meta.get("city") or "—"
    address = str(_safe(address, "—"))

    currency = str(meta.get("currency") or "EUR")

    # Facade area for per-m²
    facade_area = inputs.get("facade_area_m2")
    facade_area = float(facade_area) if isinstance(facade_area, (int, float)) else None

    def per_m2(x):
        if facade_area and isinstance(x, (int, float)):
            return x / facade_area
        return None

    # Instantaneous
    wall_w = comps.get("wall", {}).get("instantaneous_watts")
    win_w = comps.get("window", {}).get("instantaneous_watts")
    door_w = comps.get("door", {}).get("instantaneous_watts")
    total_w = totals.get("instantaneous_watts")

    # Annual (ΔT)
    wall_kwh = comps.get("wall", {}).get("annual_kwh_delta")
    win_kwh = comps.get("window", {}).get("annual_kwh_delta")
    door_kwh = comps.get("door", {}).get("annual_kwh_delta")
    total_kwh = totals.get("annual_kwh_delta")

    total_w_m2 = per_m2(total_w)
    total_kwh_m2 = per_m2(total_kwh)

    # Multi-year projection
    proj = (totals.get("multi_year_costs_delta") or {})

    # Images (required intent: RGB + overlay)
    rgb_bytes, overlay_bytes = _extract_images_from_payload(raw)

    # Build document
    doc = Document()
    _setup_document(doc)

    # Cover section
    _add_cover_page(doc, address=address, created_at=created_at, analysis_id=analysis_id)

    # Main section
    _start_main_section(doc)

    # ----------------------------
    # SUMMARY
    # ----------------------------
    _add_h1(doc, "SUMMARY (EXECUTIVE)")
    _add_par(doc, f"Location: {city}")
    _add_par(doc, f"Created: {created_at}")
    _add_par(doc, f"Analysis ID: {analysis_id}")
    doc.add_paragraph("")

    est_kwh = headline.get("estimated_annual_heat_loss_kwh")
    est_cost = headline.get("estimated_annual_cost_eur")
    confidence = headline.get("confidence")
    key_driver = headline.get("key_driver")

    _add_par(doc, f"Estimated annual heat loss (indicative): {_fmt(est_kwh, ' kWh/year', 0)}")
    _add_par(doc, f"Estimated annual cost impact (indicative): {_fmt(est_cost, f' {currency}/year', 0)}")
    _add_par(doc, f"Confidence: {_safe(confidence)}")
    if key_driver:
        _add_par(doc, f"Key driver: {key_driver}")

    doc.add_paragraph("")

    # Optional illustrative graphic (keep safe wording)
    if EER_PATH.exists():
        _add_h2(doc, "Illustrative indicator (not an EPC)")
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(EER_PATH), width=Inches(4.8))
            _add_par(doc, "This indicator is illustrative only and does not constitute an official EPC rating.")
        except Exception:
            _add_par(doc, "Illustrative indicator could not be rendered.")
        doc.add_paragraph("")

    # ----------------------------
    # TABLES
    # ----------------------------
    _add_h1(doc, "THERMAL ANALYSIS OUTPUTS")

    hourly_rows = [
        ("Wall’s heat loss at time of measurement", _fmt(wall_w, " W", 0)),
        ("Windows’ heat loss at time of measurement", _fmt(win_w, " W", 0)),
        ("Door’s heat loss at time of measurement", _fmt(door_w, " W", 0)),
        ("Façade heat loss (total)", _fmt(total_w, " W", 0)),
        ("Façade heat loss per m²", _fmt(total_w_m2, " W/m²", 2)),
    ]
    _add_legacy_style_table(
        doc,
        "HEAT LOSS AT THE TIME OF MEASUREMENT (INDICATIVE)",
        "DATA",
        hourly_rows,
    )

    annual_rows = [
        ("Annual heat loss from wall (ΔT method)", _fmt(wall_kwh, " kWh/year", 0)),
        ("Annual heat loss from windows (ΔT method)", _fmt(win_kwh, " kWh/year", 0)),
        ("Annual heat loss from doors (ΔT method)", _fmt(door_kwh, " kWh/year", 0)),
        ("Annual heat loss from façade (ΔT method)", _fmt(total_kwh, " kWh/year", 0)),
        ("Annual heat loss from façade per m²", _fmt(total_kwh_m2, " kWh/m²·year", 2)),
    ]
    _add_legacy_style_table(
        doc,
        "ANNUAL ESTIMATES OF HEAT LOSSES (INDICATIVE)",
        "DATA",
        annual_rows,
    )

    if proj:
        proj_rows = [
            ("1 year", _fmt(proj.get("1_years"), f" {currency}", 0)),
            ("5 years", _fmt(proj.get("5_years"), f" {currency}", 0)),
            ("10 years", _fmt(proj.get("10_years"), f" {currency}", 0)),
            ("20 years", _fmt(proj.get("20_years"), f" {currency}", 0)),
            ("30 years", _fmt(proj.get("30_years"), f" {currency}", 0)),
        ]
        _add_legacy_style_table(doc, "MULTI-YEAR COST PROJECTION (INDICATIVE)", "DATA", proj_rows)

    # ----------------------------
    # FIGURES (RGB + OVERLAY)
    # ----------------------------
    if rgb_bytes:
        _add_figure_page_from_bytes(doc, "RGB image (input / annotated, if provided)", rgb_bytes, width_in=6.6)
    else:
        doc.add_page_break()
        _add_h2(doc, "RGB image (input)")
        _add_par(doc, "RGB image was not provided in the payload artifacts.")

    if overlay_bytes:
        _add_figure_page_from_bytes(doc, "Thermal overlay image (indicative hotspots / annotations)", overlay_bytes, width_in=6.6)
    else:
        doc.add_page_break()
        _add_h2(doc, "Thermal overlay image")
        _add_par(doc, "Thermal overlay image was not provided in the payload artifacts.")

    # ----------------------------
    # ASSUMPTIONS
    # ----------------------------
    doc.add_page_break()
    _add_h1(doc, "ASSUMPTIONS & TRANSPARENCY")
    if assumptions:
        for a in assumptions[:30]:
            nm = a.get("name", "Assumption")
            val = a.get("value", "—")
            why = a.get("why_it_matters", "")

            p = doc.add_paragraph()
            r = p.add_run(f"{nm}: {val}")
            _set_font_run(r, FONT_BODY, 10, bold=True)

            if why:
                _add_par(doc, str(why))
            doc.add_paragraph("")
    else:
        _add_par(doc, "No assumptions were provided by the backend for this analysis.")

    # ----------------------------
    # METHOD & LIMITATIONS
    # ----------------------------
    doc.add_page_break()
    _add_h1(doc, "METHOD & LIMITATIONS")
    _add_par(
        doc,
        "This document provides an indicative screening of heat loss patterns derived from thermal imaging and AI-driven "
        "image segmentation. Values are estimated based on the inputs supplied to the analysis and are sensitive to "
        "environmental conditions (wind, solar loading, emissivity, viewing angle, and camera calibration). "
        "For regulated energy performance certification, please consult a qualified assessor and follow applicable standards."
    )

    # ----------------------------
    # DISCLAIMER (list preserved)
    # ----------------------------
    doc.add_paragraph("")
    _add_h2(doc, "DISCLAIMER")
    disclaimers = report.get("disclaimer") or []
    if disclaimers:
        for d in disclaimers[:30]:
            _add_par(doc, f"• {d}")
    else:
        _add_par(doc, DISCLAIMER_FOOTER)

    # Export
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
